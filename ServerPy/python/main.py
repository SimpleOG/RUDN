import os
from concurrent import futures

import grpc
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

from pb import generator_pb2
from pb import generator_pb2_grpc

dic = {"the_code_of_the_oop_rudn": "Шифр",
       "direction_code": "Код направления",
       "name_of_the_program": "Наименование программы",
       "block": "Блок",
       "component": "Компонента",
       "n_v_rup": "№ в РУП",
       "dop_info": "доп.инфо",
       "name_of_the_discipline_or_type_of_academic_work": "Наименование дисциплины или вида учебной работы",
       "semester_or_module": "Семестр ; Модуль",
       "weeks_per_semester_module": "Недель в семестре (модуле)",
       "type_of_educational_work": "Вид учебной работы",
       "type_of_pa_or_gia": "Вид ПА или ГИА",
       "kw_course_works": "Курс. работы",
       "kw_course_projects": "Курс. проекты",
       "course_uch_ave_ze_on_rup": "Уч. пр. (ЗЕ по РУП)",
       "pr_ze_on_rup": "Пр. пр. (ЗЕ по РУП)",
       "nir_ze_by_rup": "НИР (ЗЕ по РУП)",
       "code": "Код",
       "group_number": "Номер группы",
       "group_name":"Полное название группы",
       "of_groups": "Подгрупп",
       "subgroups": "Групп",
       "total_people": "Всего",
       "rf": "РФ",
       "foreign": "ИН",
       "standard": "Норматив",
       "calculated": "Рассчетных",
       "pk": "ПК",
       "department": "Кафедра/департамент",
       "post": "должность",
       "terms_of_attraction": "условия привлечения ",
       "full_name": "Фамилия И.О.  преподавателя",
       "a_special_feature": "Особый признак",
       "lectures": "Лекции",
       "practice_or_seminars": "Практика / Семинары",
       "lab_works_or_clinical_classes": "Лаб. работы / Клинические занятия",
       "current_control": "Текущий контроль",
       "interim_certification_po_for_brs": "Промежуточная аттестация (ПА) по БРС",
       "registration_of_pa_results": "Оформление результатов ПА",
       "ongoing_consultations_on_the_discipline": "Текущие консультации по дисциплине",
       "course_works": "Курсовые работы",
       "course_projects": "Курсовые проекты",
       "educational_practice": "Учебная практика",
       "proc_pedagogical_and_pre_graduate_practices": "Произв. педагогическая и преддипломная практики",
       "nir": "НИР",
       "practices_including_research_of_digital_magistracies": "Практики (в т.ч. НИР) цифровых магистратур",
       "reviewing_the_abstracts_of_graduate_students": "Рецензирование рефератов аспирантов",
       "candidates_exam": "Кандидатский экзамен",
       "scientific_guidance": "Научное руководство",
       "the_leadership_of_the_wrc_or_the_nkr": "Руководство ВКР или НКР в том числе Организация и сопровождение Первичной аккредитации МИ",
       "review_of_the_wrc": "Рецензирование ВКР",
       "gek": "ГЭК ",
       "total": "ИТОГО"}


class FileGeneratorServicer(generator_pb2_grpc.FileGeneratorServicer):
    def Generate(self, request, context):
        print("Генерация начата")
        filepath = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '../ForDownload'))
        name = request.name
        filepath = os.path.join(filepath, name + ".doc")
        document = Document()
        document.add_heading(request.name, 0)
        sections = document.sections
        for section in sections:
         section.left_margin = Inches(1.0)  # Левое поле
         section.right_margin = Inches(1.0)  # Правое поле
         section.top_margin = Inches(1.0)  # Верхнее поле
         section.bottom_margin = Inches(1.0)  # Нижнее поле
        table = document.add_table(rows=1, cols=len(list(request.data)[0].map.keys()))
        table.style = 'Table Grid'
        for i, key in enumerate(list(request.data)[0].map.keys()):
            cell = table.cell(0, i)
            cell.text = dic[key]
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for my_map in request.data:
            count=0
            cells = table.add_row().cells
            for i in my_map.map.values():
                cells[count].text = str(i)
                count+=1
        document.save(filepath)
        response = generator_pb2.GenerateResponse(
            filepath=filepath,
            status="Success",
        )
        return response


def serve():
    server = grpc.server(futures.ThreadPoolExecutor(max_workers=10))
    generator_pb2_grpc.add_FileGeneratorServicer_to_server(FileGeneratorServicer(), server)
    server.add_insecure_port('[::]:1111')
    print(f"started ")
    server.start()
    server.wait_for_termination()


if __name__ == '__main__':
    serve()
